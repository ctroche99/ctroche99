import pathlib
import requests
import json
import docx
from datetime import datetime
from datetime import date
from sys import exit
import pytictoc
import tkinter
import tkinter.messagebox
from tkinter import ttk as GUI
from tkinter import E,W
from tkinter import filedialog

def initDirectories(FileName: str) -> pathlib.Path: #check if file exists and if not, returns file path to newly made file
    curDir = str(pathlib.Path(__file__).parent.resolve())
    FilePath = curDir+FileName
    FilePath = pathlib.Path(FilePath)
    if FilePath.exists() == False:
        dirOut_file = open(str(FilePath),'w')
        dirOut_file.writelines(curDir+"\n")
        dirOut_file.writelines(curDir+"\n")
        dirOut_file.close()
    return FilePath
    
class appHead:
    def __init__(self, OA: dict) -> None:
        self.appNum :str = OA['patentApplicationNumber'][0]
        self.appNum = self.appNum[0:2]+'/'+self.appNum[2:5]+','+self.appNum[5:]
        self.matter :str = OA['applicantFileReference'][0]
        self.claims :str = str(OA['effectiveClaimTotalQuantity'])
        self.examNum :str = OA['examinerEmployeeNumber'][0]
        self.title :str = OA['inventionTitle'][0]
        self.custNum :str = str(OA['customerNumber'])
        self.artUnit :str = str(OA['groupArtUnitNumber'])
        self.confNum :str = str(OA['patentApplicationConfirmationNumber'])
        self.filDate :str = OA['filingDate'][0:10]
        self.OAdate :str = OA['submissionDate'][0:10]
        self.text :str = OA['bodyText'][0]
        self.curDate :str = date.today().strftime('%B')+' '+date.today().strftime('%d')+', '+date.today().strftime('%Y')

class rejData:
    def __init__(self) -> None:
        self.entry112 :bool= False #Flags for each rejection type to reduce memory usage
        self.entry101 :bool= False 
        self.entry102 :bool= False
        self.entry103 :bool= False
        self.entryDbl :bool= False
        self.entryAlw :bool= False
        self.entryObj :bool = False
        self.rej_112 :list = [] #Lists to store rejections parsed from OA body text
        self.rej_101 :list = []
        self.rej_102 :list = []
        self.rej_103 :list = []
        self.rej_Dbl :list = []
        self.obj_Alw :list = []

class rej:
    def __init__(self,text: str,rejection: str):
        catRej_Flag = text.find("as applied to")
        if catRej_Flag != -1:
            self.text = catRej(text, catRej_Flag)
        else:
            self.text = text
        self.type = rejection
        self.references = getReferences(text) #Technically not necessary, leave here anyways
        self.claims = getClaims(text) #functioning for now...

def errorMsg(errorStr: str):
    tkinter.messagebox.showinfo("ERROR", errorStr)
    exit()

def catRej(text: str, index: int) -> str:
    splitPoint = text.find("as obvious over")
    return text[0:(splitPoint+15)]+" the references "+text[index:]

def getReferences(text: str):
    pass

def getRejectionType(text: str):
    pass #may need to implement later, but right now string passing works

def getClaims(text: str) -> str: 
    claimStart = text.find("Claim")
    if claimStart != -1: #if word Claim is not found, skip
        claims = []
        while(text[claimStart].isspace() == False):#account for different variations of word Claim, Claims, Claim(s)
            claimStart += 1
        i = claimStart
        while i < len(text): #parse out numbers and common groups of numbers for claims
            if text[i] == ' ' or text[i] == ',' or text[i] == 'a' or text[i] == 'n' or text[i] == 'd':
                i += 1
                continue
            if text[i].isdigit():
                claimBuf = ''
                for j in range(i,i+6): #realistically claims will not excedd 99 or 99-99
                    if text[j].isdigit() or text[j] == '-':
                        claimBuf = claimBuf + text[j]
                    else: #Get out of loop once end of numeric combination is found
                        claims.append(claimBuf)
                        i = j #avoid copying in numbers already gone through
                        break
            else: #End once all claims are found
                break
            i += 1
    else: #handler for not finding any claims
        #print("No claims found.")
        claims = ["N/A"]
    #print(claims,"found in this rejection")
    return claims

def pullUSPTO(input: str) -> dict:
    if input.isdigit() == True:
        criteria = 'patentApplicationNumber%3A%20'+input
    else:
        errorMsg('Invalid Input. Please enter Application Serial Number without commas or slashes.')
        #criteria = r"applicantFileReference%3A%20%22"+userIn+"%22" Enables Matter Search feature
    
    url = 'https://developer.uspto.gov/ds-api/oa_actions/v1/records'
    headers = requests.structures.CaseInsensitiveDict()
    headers["Accept"] = "application/json"
    headers["Content-Type"] = "application/x-www-form-urlencoded"
    data = "criteria="+criteria+"&start=0&rows=100"

    req_Info = requests.post(url,headers=headers,data=data)

    if req_Info.status_code != 200:
        errorMsg('Unable to retrieve data from USPTO. Make sure Matter or Application number is correctly typed.')

    #with open("OA_Info.json","w") as out_file:
    req_dict = json.loads(req_Info.content)
    #   req_json = json.dumps(req_dict, indent = 4)
    #  out_file.write(req_json)
    return req_dict

def parseOA(OAtext: list) -> rejData:
    Rejections = rejData()
    Rejections.entry112 = False #Flags for each rejection type to reduce memory usage
    Rejections.entry101 = False 
    Rejections.entry102 = False
    Rejections.entry103 = False
    Rejections.entryDbl = False
    Rejections.entryAlw = False
    Rejections.entryObj = False   

    for curPara in OAtext: #parse rejections out of text
        if curPara.find("rejected under 35 U.S.C. 112") != -1 or curPara.find("rejected under 35 U.S.C. § 112") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            Rejections.rej_112.append(rej(curPara,"35 U.S.C. § 112"))
            #print(curPara,'\n')
            Rejections.entry112 = True

        elif curPara.find("rejected under 35 U.S.C. 101") != -1 or curPara.find("rejected under 35 U.S.C. § 101") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            Rejections.rej_101.append(rej(curPara,"35 U.S.C. § 101"))
            #print(curPara,'\n')
            Rejections.entry101 = True

        elif curPara.find("rejected under 35 U.S.C. 102") != -1 or curPara.find("rejected under 35 U.S.C. § 102") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            Rejections.rej_102.append(rej(curPara,"35 U.S.C. § 102"))
            #print(curPara,'\n')
            Rejections.entry102 = True

        elif curPara.find("rejected under 35 U.S.C. 103") != -1 or curPara.find("rejected under 35 U.S.C. § 103") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            Rejections.rej_103.append(rej(curPara,"35 U.S.C. § 103"))
            #print(curPara,'\n')
            Rejections.entry103 = True

        elif curPara.find("rejected on the ground of nonstatutory double patenting") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            Rejections.rej_Dbl.append(rej(curPara,"Nonstatutory Double Patenting"))
            #print(curPara,'\n')
            Rejections.entryDbl = True

        elif curPara.find("would be allowable") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            Rejections.obj_Alw.append(rej(curPara,"would be allowable"))#Assuming allowable claims are mentioned only once
            #print(curPara,'\n')
            Rejections.entryAlw = True    
    
    return Rejections

def generateWordTemplate(rejection_data: rejData, app_data: appHead,tempDir :tkinter.StringVar,outDir :tkinter.StringVar):
    tempPath = pathlib.Path(tempDir.get()+"/"+"OA_Template.docx")
    if tempPath.exists() == False:
        errorMsg('ERROR: Office Action Template not Found in given directory: '+tempDir.get())

    Template = docx.Document(tempDir.get()+"/"+"OA_Template.docx")
    TemplateFileName = (outDir.get()+"/"+app_data.matter+"_OATemplate.docx")
    Template.save(TemplateFileName) 
    
    for para in Template.paragraphs:#Find starting paragraph for insertion before
        curPara = str(para.text)
        if(curPara.find("Independent Claim 1") != -1):
            startPara = para
            break   #Generate Paragraphs

    romNums = ['II.','III.','IV.','V.','VI.','VII.']
    i = 0

    if rejection_data.entryAlw == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tAllowable Subject Mattter",style="Paragraph Header")
        i += 1
        AlwPara = startPara.insert_paragraph_before("\tApplicant acknowledges with appreciation the Examiner’s indication that claims "+" are directed to allowable subject matter.\n")
        AlwPara.paragraph_format.line_spacing = 1.5

    if rejection_data.entryObj == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tObjection to the Claims",style="Paragraph Header")
        i += 1
        ObjPara = startPara.insert_paragraph_before("\tThe Office Action objects to [OASUMMARYOBJECTED].  In effort to advance prosecution and without acquiescing to the propriety of the objection, the claims have been amended to address such objections.\n")
        ObjPara.paragraph_format.line_spacing = 1.5

    if rejection_data.entryDbl == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tDouble Patenting",style="Paragraph Header")
        i += 1
        DblPara = startPara.insert_paragraph_before("\tWithout acquiescing to propriety of the rejection, Applicant will consider submitting a terminal disclaimer, if such disclaimer is in fact needed, to overcome the double patenting rejections at a time when the claims are otherwise in condition for allowance.\n")
        DblPara.paragraph_format.line_spacing = 1.5

    if rejection_data.entry112 == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tRejection under 35 U.S.C. §§ 112",style="Paragraph Header")
        i += 1
        text112 = ''
        for rej in rejection_data.rej_112:#Combine each rejection into one
            text112 = text112+rej.text+' '
        Para112 = startPara.insert_paragraph_before("\tThe Examiner has rejected claims "+"under 35 U.S.C. § 112 for failing to particularly point out and distinctly claim the subject matter which the inventor regards as the invention.  Without acquiescing to the propriety of the rejection, Applicant has amended the claims to remedy the deficiencies.  Accordingly, Applicant respectfully requests withdrawal of the § 112 rejections\n")
        Para112.paragraph_format.line_spacing = 1.5

    if rejection_data.entry101 == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tRejections under 35 U.S.C. §§ 101",style="Paragraph Header")
        i += 1
        text101 = ''
        for rej in rejection_data.rej_101:#Combine each rejection into one
            text101 = text101+rej.text+' '
        Para101 = startPara.insert_paragraph_before("\t"+text101+"\n")
        Para101.paragraph_format.line_spacing = 1.5  

    if rejection_data.entry102 == True or rejection_data.entry103 == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tRejections under 35 U.S.C. §§ 102 and 103",style="Paragraph Header")
        i += 1
        text102_3 = ''
        for rej in rejection_data.rej_103:#Combine each rejection into one
            text102_3 = text102_3+rej.text+' '
        Para102_3 = startPara.insert_paragraph_before("\t"+text102_3+"\n")  
        Para102_3.paragraph_format.line_spacing = 1.5   

    #Fills Table on front page
    Template.tables[0]._cells[2].text = 'Inventor' #Inventor
    Template.tables[0]._cells[5].text = 'Examiner' #Examiner
    Template.tables[0]._cells[14].text = app_data.appNum #Serial Number
    Template.tables[0]._cells[17].text = app_data.artUnit #Group Art Unit
    Template.tables[0]._cells[26].text = app_data.filDate #Filing Date
    Template.tables[0]._cells[29].text = app_data.custNum #Customer Number
    Template.tables[0]._cells[38].text = app_data.confNum #Confirmation Number
    Template.tables[0]._cells[50].text = app_data.title #Title
    Template.tables[0]._cells[53].text = app_data.matter #Matter Number

    #Fill small paragraph on first page
    OAdate = datetime.fromisoformat(app_data.OAdate)
    Template.paragraphs[11].text = '\tIn response to the Non-Final Office Action dated '+OAdate.strftime('%B')+' '+OAdate.strftime('%d')+', '+OAdate.strftime('%Y')+', please amend the present application as set forth below and consider the remarks that follow.'

    #Fill in Header
    TempHeader = Template.sections[0].header
    TempHeader.paragraphs[0].text = 'Serial Number: '+app_data.appNum+'\t\tAttorney Docket Number: '+app_data.matter
    TempHeader.paragraphs[1].text = 'Response Dated: '+app_data.curDate
    TempHeader.paragraphs[2].text = 'Office Action Date: '+app_data.OAdate

    Template.save(TemplateFileName)

class mainWin:
    def __init__(self,root: tkinter.Tk,tempDirectory :str, outDirectory :str,dirPath :pathlib.Path) -> None:
        self.tempDir = tkinter.StringVar(value=tempDirectory)
        self.outDir = tkinter.StringVar(value=outDirectory)
        self.dirPath = dirPath

        self.root = root
        self.root.title('Template Maker v2.1')
        self.root.columnconfigure(0,weight=1)
        self.root.rowconfigure(0,weight=1)

        dirPath_str = str(dirPath)
        dirPath_index = dirPath_str.find('TemplateMaker_Directories.txt')
        imagePath_str = dirPath_str[:dirPath_index]+"DANDM.ico"
        imagePath = pathlib.Path(imagePath_str)
        if imagePath.exists() == True:
            root.iconbitmap(imagePath)

        self.frame = GUI.Frame(self.root)
        self.frame.grid(column=0,row=0)

        self.input = tkinter.StringVar()
        self.entry = GUI.Entry(self.frame, width=20,textvariable=self.input)
        self.entry.grid(column=2,row=1,columnspan=2,sticky=W)
        self.entry.grid_configure(padx=5,pady=5)

        self.entryLabel = GUI.Label(self.frame,text='Enter Application #:')
        self.entryLabel.grid(column=1,row=1)
        self.entryLabel.grid_configure(padx=5,pady=10)

        self.optButn = GUI.Button(self.frame, text='Options...', command=self.makeOptWin)
        self.optButn.grid(column=1, row=2, sticky=W)
        self.optButn.grid_configure(padx=3,pady=10)

        self.entButn = GUI.Button(self.frame, text='Enter', command=self.makeTemp)
        self.entButn.grid(column=2, row=2)
        self.entButn.grid_configure(padx=3,pady=10)

        self.exButn = GUI.Button(self.frame, text='Cancel', command=exit)
        self.exButn.grid(column=3,row=2)
        self.exButn.grid_configure(padx=3,pady=10)
        self.entry.focus()
        self.root.bind("<Return>",self.makeTemp)

    def makeOptWin(self):
        self.optWin = tkinter.Toplevel(self.root)
        self.optApp = optWin(self.root,self.optWin,self.tempDir,self.outDir,self.dirPath)
    
    def makeTemp(self,event=None):
        userIn = self.input.get()
        time = pytictoc.TicToc()
        #userIn = input("Please enter Application Number or Matter Number without commas or slashes: ")
        time.tic()

        req_dict = pullUSPTO(userIn)

        if req_dict['response']['numFound'] == 0:
            errorMsg('Unable to retrieve data from USPTO. Make sure Matter or Application number is correctly typed.')

        print("Application Found.")
        subDate = []
        for result in req_dict['response']['docs']:#Find latest Office Action, this may be avoided by just the last entry returned...
            subDate.append(result['submissionDate'])#Need to investigate ^^
        subDatesort = sorted(subDate)
        dateIndex = subDatesort[-1]
        OA_Index = subDate.index(dateIndex)

        OA_doc = req_dict['response']['docs'][OA_Index]
        OA = appHead(OA_doc)
        print('Generating Template for Office Action dated '+OA_doc['submissionDate'][0:10])
        print('Application Number: '+OA.appNum)
        print('Matter: '+OA.matter)
        print('Title: '+OA.title)
        print('Filed: '+OA.filDate)
        print("Claims: "+str(OA.claims))

        OA_text_list = OA.text.split('\n')
        OArejData = parseOA(OA_text_list)

        generateWordTemplate(OArejData, OA,self.tempDir,self.outDir)

        print('End of Program.')
        time.toc()
        tkinter.messagebox.showinfo("Generation Complete.", "Template Successfully Generated.")

class optWin:
    def __init__(self,master :tkinter.Tk, root :tkinter.Tk,tempDir :tkinter.StringVar,outDir :tkinter.StringVar,dirPath: pathlib.Path):
        
        self.root = root#Option Window Root
        self.frame = GUI.Frame(self.root)
        self.frame.grid(column=0,row=0)
        self.root.transient(master)#Original Window's Root

        self.tempDirLabel = GUI.Label(self.frame,text='Template File Directory:')
        self.tempDirLabel.grid(column=1,row=1,sticky=W)
        self.tempDirLabel.grid_configure(padx=5,pady=1)
    
        self.tempDirEntry = GUI.Entry(self.frame, width=70, textvariable=tempDir)
        self.tempDirEntry.grid(column=1,row=2,columnspan=10)
        self.tempDirEntry.grid_configure(padx=5,pady=5)

        self.tempDirBut = GUI.Button(self.frame, text='...',width=3, command=lambda: self.tempBrowse(tempDir))
        self.tempDirBut.grid(column=11,row=2,sticky=W)
        self.tempDirBut.grid_configure(padx=3,pady=7)

        self.outDirLabel = GUI.Label(self.frame,text='File Output Directory:')
        self.outDirLabel.grid(column=1,row=3,sticky=W)
        self.outDirLabel.grid_configure(padx=5,pady=1)
    
        self.outDirEntry = GUI.Entry(self.frame, width=70, textvariable=outDir)
        self.outDirEntry.grid(column=1,row=4,columnspan=10)
        self.outDirEntry.grid_configure(padx=5,pady=5)

        self.outDirBut = GUI.Button(self.frame, text='...',width=3,command=lambda: self.outBrowse(outDir))
        self.outDirBut.grid(column=11,row=4,sticky=W)
        self.outDirBut.grid_configure(padx=3,pady=7)

        self.entButn = GUI.Button(self.frame, text='Enter', command=lambda: self.entPress(tempDir,outDir,dirPath))
        self.entButn.grid(column=10, row=5,sticky=E)
        self.entButn.grid_configure(padx=3,pady=10)

        self.exButn = GUI.Button(self.frame, text='Cancel', command=self.root.destroy)
        self.exButn.grid(column=11,row=5,sticky=W)
        self.exButn.grid_configure(padx=3,pady=10)

    def tempBrowse(self,directory: tkinter.StringVar):
        filename = filedialog.askdirectory(initialdir='/',title='Select Folder')
        directory.set(filename)

    def outBrowse(self, directory: tkinter.StringVar):
        filename = filedialog.askdirectory(initialdir='/',title='Select Folder')
        directory.set(filename)

    def entPress(self,tempDir :tkinter.StringVar,outDir :tkinter.StringVar,dirPath :pathlib.Path):
        dirOut_file = open(str(dirPath),'w') #write the newly added directories to default file
        dirOut_file.write(tempDir.get()+'\n')
        dirOut_file.write(outDir.get()+'\n')
        dirOut_file.close()
        self.root.destroy()

def main():
    directoryFilePath = initDirectories('\TemplateMaker_Directories.txt')
    dirIn_file = open(str(directoryFilePath),'r')
    templateDirectory = dirIn_file.readline()
    templateDirectory = templateDirectory[:-1]
    outputDirectory = dirIn_file.readline()
    outputDirectory = outputDirectory[:-1]
    dirIn_file.close()

    root = tkinter.Tk()
    mainWindow = mainWin(root,templateDirectory,outputDirectory,directoryFilePath)
    root.mainloop()

    
if __name__ == '__main__':
    main()