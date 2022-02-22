import docx
import sys

#Incorporate USPTO API interfacing for information retrieval
#Returned bodies do not solve citation parsing problems
def GetHeader(appNum: str):
    pass

def catRej(text: str, index: int) -> str:
    splitPoint = text.find("as obvious over")
    return text[0:(splitPoint+15)]+" the references "+text[index:]

def getReferences(text: str):
    pass

def getRejectionType(text: str):
    pass #may need to implement later, but right now string passing works

def getClaims(text: str) -> str: 
    claimStart = text.find("Claim")
    if claimStart != -1: #if word Claim is not found, skip
        claims = []
        while(text[claimStart].isspace() == False):#account for different variations of word Claim, Claims, Claim(s)
            claimStart += 1
        i = claimStart
        while i < len(text): #parse out numbers and common groups of numbers for claims
            if text[i] == ' ' or text[i] == ',' or text[i] == 'a' or text[i] == 'n' or text[i] == 'd':
                i += 1
                continue
            if text[i].isdigit():
                claimBuf = ''
                for j in range(i,i+6): #realistically claims will not excedd 99 or 99-99
                    if text[j].isdigit() or text[j] == '-':
                        claimBuf = claimBuf + text[j]
                    else: #Get out of loop once end of numeric combination is found
                        claims.append(claimBuf)
                        i = j #avoid copying in numbers already gone through
                        break
            else: #End once all claims are found
                break
            i += 1
    else: #handler for not finding any claims
        print("No claims found.")
        claims = ["N/A"]
    print(claims,"found in this rejection")
    return claims

class appHead:
    pass

class Rejection:
    def __init__(self,text: str,rejection: str):
        catRej_Flag = text.find("as applied to")
        if catRej_Flag != -1:
            self.text = catRej(text, catRej_Flag)
        else:
            self.text = text
        self.type = rejection
        self.references = getReferences(text) #Technically not necessary, leave here anyways
        self.claims = getClaims(text) #functioning for now...
        
        

if len(sys.argv) != 2:#enable user input file names
    #incorporate PDF support

    #potential way to simplify PATH used?
    print(r"./TemplateMaker [OfficeAction.docx OR *.pdf]")
    print("\n")
    exit() 
    
    #change to documents or downloads folder. somewhere generic.

if sys.argv[1].find(".docx") != -1: 
    OA = docx.Document("C:/Users/ctroc/Documents/Python/TemplateMaker/"+sys.argv[1])
    entry112 = False #Flags for each rejection type to reduce memory usage
    entry101 = False 
    entry102 = False
    entry103 = False
    entryDbl = False
    entryAlw = False
    entryObj = False    
    
    #GetHeader(APPNUM)
    for para in OA.paragraphs: #parse rejections out of .docx
        curPara = str(para.text)
        if curPara.find("rejected under 35 U.S.C. 112") != -1 or curPara.find("rejected under 35 U.S.C. § 112") != -1:
            if entry112 == False:
                rej_112 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_112.append(Rejection(curPara,"35 U.S.C. § 112"))
            print(curPara,'\n')
            entry112 = True

        elif curPara.find("rejected under 35 U.S.C. 101") != -1 or curPara.find("rejected under 35 U.S.C. § 101") != -1:
            if entry101 == False:
                rej_101 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_101.append(Rejection(curPara,"35 U.S.C. § 101"))
            print(curPara,'\n')
            entry101 = True

        elif curPara.find("rejected under 35 U.S.C. 102") != -1 or curPara.find("rejected under 35 U.S.C. § 102") != -1:
            if entry102 == False:
                rej_102 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_102.append(Rejection(curPara,"35 U.S.C. § 102"))
            print(curPara,'\n')
            entry102 = True

        elif curPara.find("rejected under 35 U.S.C. 103") != -1 or curPara.find("rejected under 35 U.S.C. § 103") != -1:
            if entry103 == False:
                rej_103 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_103.append(Rejection(curPara,"35 U.S.C. § 103"))
            print(curPara,'\n')
            entry103 = True

        elif curPara.find("rejected on the ground of nonstatutory double patenting") != -1:
            if entryDbl == False:
                rej_dbl = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_dbl.append(Rejection(curPara,"Nonstatutory Double Patenting"))
            print(curPara,'\n')
            entryDbl = True

        elif curPara.find("would be allowable") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            obj_Alw = Rejection(curPara,"would be allowable")#Assuming allowable claims are mentioned only once
            print(curPara,'\n')
            entryAlw = True 

elif sys.argv[1].find(".pdf") != -1:
    #incorporate PDF usage maybe, API might be superior option
    pass
elif sys.argv[1].isnumeric == True:
    userIn = input("Is this an application number?\n")
    if userIn != 'Y':
        print("Operation Cancelled.\n")
        exit()
    #incorporate USPTO API for pulling OA Data

#Write out to the template document.
Template = docx.Document("C:/Users/ctroc/Documents/Python/TemplateMaker/OA_Template.docx")
Template.save("C:/Users/ctroc/Documents/Python/TemplateMaker/Test_Template.docx")   
for para in Template.paragraphs:#Find starting paragraph for insertion before
    curPara = str(para.text)
    if(curPara.find("Independent Claim 1") != -1):
        startPara = para
        break   #Generate Paragraphs

romNums = ['II.','III.','IV.','V.','VI.','VII.']
i = 0

if entryAlw == True:
    startPara.insert_paragraph_before("\t"+romNums[i]+" Allowable Subject Mattter",style="Paragraph Header")
    i += 1
    AlwPara = startPara.insert_paragraph_before("\tApplicant acknowledges with appreciation the Examiner’s indication that claims "+" are directed to allowable subject matter.\n")
    AlwPara.paragraph_format.line_spacing = 1.5

if entryObj == True:
    startPara.insert_paragraph_before("\t"+romNums[i]+" Objection to the Claims",style="Paragraph Header")
    i += 1
    ObjPara = startPara.insert_paragraph_before("\tThe Office Action objects to [OASUMMARYOBJECTED].  In effort to advance prosecution and without acquiescing to the propriety of the objection, the claims have been amended to address such objections.\n")
    ObjPara.paragraph_format.line_spacing = 1.5

if entryDbl == True:
    startPara.insert_paragraph_before("\t"+romNums[i]+" Double Patenting",style="Paragraph Header")
    i += 1
    DblPara = startPara.insert_paragraph_before("\tWithout acquiescing to propriety of the rejection, Applicant will consider submitting a terminal disclaimer, if such disclaimer is in fact needed, to overcome the double patenting rejections at a time when the claims are otherwise in condition for allowance.\n")
    DblPara.paragraph_format.line_spacing = 1.5

if entry112 == True:
    startPara.insert_paragraph_before("\t"+romNums[i]+" Rejection under 35 U.S.C. §§ 112",style="Paragraph Header")
    i += 1
    text112 = ''
    for Rej in rej_112:#Combine each rejection into one
        text112 = text112+Rej.text+' '
    Para112 = startPara.insert_paragraph_before("\tThe Examiner has rejected claims "+"under 35 U.S.C. § 112 for failing to particularly point out and distinctly claim the subject matter which the inventor regards as the invention.  Without acquiescing to the propriety of the rejection, Applicant has amended the claims to remedy the deficiencies.  Accordingly, Applicant respectfully requests withdrawal of the § 112 rejections\n")
    Para112.paragraph_format.line_spacing = 1.5

if entry101 == True:
    startPara.insert_paragraph_before("\t"+romNums[i]+" Rejections under 35 U.S.C. §§ 101",style="Paragraph Header")
    i += 1
    text101 = ''
    for Rej in rej_101:#Combine each rejection into one
        text101 = text101+Rej.text+' '
    Para101 = startPara.insert_paragraph_before("\t"+text101+"\n")
    Para101.paragraph_format.line_spacing = 1.5  

if entry102 == True or entry103 == True:
    startPara.insert_paragraph_before("\t"+romNums[i]+" Rejections under 35 U.S.C. §§ 102 and 103",style="Paragraph Header")
    i += 1
    text102_3 = ''
    for Rej in rej_103:#Combine each rejection into one
        text102_3 = text102_3+Rej.text+' '
    Para102_3 = startPara.insert_paragraph_before("\t"+text102_3+"\n")  
    Para102_3.paragraph_format.line_spacing = 1.5   

Template.save("C:/Users/ctroc/Documents/Python/TemplateMaker/Test_Template.docx")
print("End of Program")

