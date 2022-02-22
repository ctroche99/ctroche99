import requests
import json
import docx
import tkinter
import tkinter.messagebox
from tkinter import ttk as GUI
from datetime import date
import pytictoc

def errorMsg(errorStr: str):
    tkinter.messagebox.showinfo("ERROR", errorStr)
    exit()

def catRej(text: str, index: int) -> str:
    splitPoint = text.find("as obvious over")
    return text[0:(splitPoint+15)]+" the references "+text[index:]

def getReferences(text: str):
    pass

def getRejectionType(text: str):
    pass #may need to implement later, but right now string passing works

def getClaims(text: str) -> str: 
    claimStart = text.find("Claim")
    if claimStart != -1: #if word Claim is not found, skip
        claims = []
        while(text[claimStart].isspace() == False):#account for different variations of word Claim, Claims, Claim(s)
            claimStart += 1
        i = claimStart
        while i < len(text): #parse out numbers and common groups of numbers for claims
            if text[i] == ' ' or text[i] == ',' or text[i] == 'a' or text[i] == 'n' or text[i] == 'd':
                i += 1
                continue
            if text[i].isdigit():
                claimBuf = ''
                for j in range(i,i+6): #realistically claims will not excedd 99 or 99-99
                    if text[j].isdigit() or text[j] == '-':
                        claimBuf = claimBuf + text[j]
                    else: #Get out of loop once end of numeric combination is found
                        claims.append(claimBuf)
                        i = j #avoid copying in numbers already gone through
                        break
            else: #End once all claims are found
                break
            i += 1
    else: #handler for not finding any claims
        #print("No claims found.")
        claims = ["N/A"]
    #print(claims,"found in this rejection")
    return claims

class appHead:
    def __init__(self, OA: dict) -> None:
        self.appNum :str = OA['patentApplicationNumber'][0]
        self.appNum = self.appNum[0:2]+'/'+self.appNum[2:5]+','+self.appNum[5:]
        self.matter :str = OA['applicantFileReference'][0]
        self.claims :str = str(OA['effectiveClaimTotalQuantity'])
        self.examNum :str = OA['examinerEmployeeNumber'][0]
        self.title :str = OA['inventionTitle'][0]
        self.custNum :str = str(OA['customerNumber'])
        self.artUnit :str = str(OA['groupArtUnitNumber'])
        self.confNum :str = str(OA['patentApplicationConfirmationNumber'])
        self.filDate :str = OA['filingDate'][0:10]
        self.OAdate :str = OA['submissionDate'][0:10]
        self.text :str = OA['bodyText'][0]
        self.curDate :str = date.today().strftime('%B')+' '+date.today().strftime('%d')+', '+date.today().strftime('%Y')

class Rejection:
    def __init__(self,text: str,rejection: str):
        catRej_Flag = text.find("as applied to")
        if catRej_Flag != -1:
            self.text = catRej(text, catRej_Flag)
        else:
            self.text = text
        self.type = rejection
        self.references = getReferences(text) #Technically not necessary, leave here anyways
        self.claims = getClaims(text) #functioning for now...

def makeTemplate(*args):
    userIn = input.get()
    time = pytictoc.TicToc()
    #userIn = input("Please enter Application Number or Matter Number without commas or slashes: ")
    time.tic()
    if userIn.isdigit() == True:
        criteria = 'patentApplicationNumber%3A%20'+userIn
    else:
        errorMsg('Invalid Input. Please enter Application Serial Number without commas or slashes.')
        #criteria = r"applicantFileReference%3A%20%22"+userIn+"%22"

    url = 'https://developer.uspto.gov/ds-api/oa_actions/v1/records'
    headers = requests.structures.CaseInsensitiveDict()
    headers["Accept"] = "application/json"
    headers["Content-Type"] = "application/x-www-form-urlencoded"
    data = "criteria="+criteria+"&start=0&rows=100"

    req_Info = requests.post(url,headers=headers,data=data)

    if req_Info.status_code != 200:
        errorMsg('Unable to retrieve data from USPTO. Make sure Matter or Application number is correctly typed.')

    #with open("OA_Info.json","w") as out_file:
    req_dict = json.loads(req_Info.content)
    #   req_json = json.dumps(req_dict, indent = 4)
    #  out_file.write(req_json)

    if req_dict['response']['numFound'] == 0:
        errorMsg('Unable to retrieve data from USPTO. Make sure Matter or Application number is correctly typed.')

    print("Application Found.")
    subDate = []
    for result in req_dict['response']['docs']:#Find latest Office Action, this may be avoided by just the last entry returned...
        subDate.append(result['submissionDate'])#Need to investigate ^^
    subDatesort = sorted(subDate)
    dateIndex = subDatesort[-1]
    OA_Index = subDate.index(dateIndex)

    OA_doc = req_dict['response']['docs'][OA_Index]
    OA = appHead(OA_doc)
    print('Generating Template for Office Action dated '+OA_doc['submissionDate'][0:10])
    print('Application Number: '+OA.appNum)
    print('Matter: '+OA.matter)
    print('Title: '+OA.title)
    print('Filed: '+OA.filDate)
    print("Claims: "+str(OA.claims))

    OA_text_list = OA.text.split('\n')
    entry112 = False #Flags for each rejection type to reduce memory usage
    entry101 = False 
    entry102 = False
    entry103 = False
    entryDbl = False
    entryAlw = False
    entryObj = False    

    for curPara in OA_text_list: #parse rejections out of text
        if curPara.find("rejected under 35 U.S.C. 112") != -1 or curPara.find("rejected under 35 U.S.C. § 112") != -1:
            if entry112 == False:
                rej_112 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_112.append(Rejection(curPara,"35 U.S.C. § 112"))
            #print(curPara,'\n')
            entry112 = True

        elif curPara.find("rejected under 35 U.S.C. 101") != -1 or curPara.find("rejected under 35 U.S.C. § 101") != -1:
            if entry101 == False:
                rej_101 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_101.append(Rejection(curPara,"35 U.S.C. § 101"))
            #print(curPara,'\n')
            entry101 = True

        elif curPara.find("rejected under 35 U.S.C. 102") != -1 or curPara.find("rejected under 35 U.S.C. § 102") != -1:
            if entry102 == False:
                rej_102 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_102.append(Rejection(curPara,"35 U.S.C. § 102"))
            #print(curPara,'\n')
            entry102 = True

        elif curPara.find("rejected under 35 U.S.C. 103") != -1 or curPara.find("rejected under 35 U.S.C. § 103") != -1:
            if entry103 == False:
                rej_103 = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_103.append(Rejection(curPara,"35 U.S.C. § 103"))
            #print(curPara,'\n')
            entry103 = True

        elif curPara.find("rejected on the ground of nonstatutory double patenting") != -1:
            if entryDbl == False:
                rej_dbl = []
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            rej_dbl.append(Rejection(curPara,"Nonstatutory Double Patenting"))
            #print(curPara,'\n')
            entryDbl = True

        elif curPara.find("would be allowable") != -1:
            curParaIndex = curPara.find('laim')
            curPara = curPara[curParaIndex-1:]
            obj_Alw = Rejection(curPara,"would be allowable")#Assuming allowable claims are mentioned only once
            #print(curPara,'\n')
            entryAlw = True 

    Template = docx.Document("C:/Users/ctroc/Documents/Python/TemplateMaker/OA_Template.docx")
    TemplateFileName = "C:/Users/ctroc/Documents/Python/TemplateMaker/"+OA.matter+"_OATemplate.docx"
    Template.save(TemplateFileName) 
    
    for para in Template.paragraphs:#Find starting paragraph for insertion before
        curPara = str(para.text)
        if(curPara.find("Independent Claim 1") != -1):
            startPara = para
            break   #Generate Paragraphs

    romNums = ['II.','III.','IV.','V.','VI.','VII.']
    i = 0

    if entryAlw == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tAllowable Subject Mattter",style="Paragraph Header")
        i += 1
        AlwPara = startPara.insert_paragraph_before("\tApplicant acknowledges with appreciation the Examiner’s indication that claims "+" are directed to allowable subject matter.\n")
        AlwPara.paragraph_format.line_spacing = 1.5

    if entryObj == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tObjection to the Claims",style="Paragraph Header")
        i += 1
        ObjPara = startPara.insert_paragraph_before("\tThe Office Action objects to [OASUMMARYOBJECTED].  In effort to advance prosecution and without acquiescing to the propriety of the objection, the claims have been amended to address such objections.\n")
        ObjPara.paragraph_format.line_spacing = 1.5

    if entryDbl == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tDouble Patenting",style="Paragraph Header")
        i += 1
        DblPara = startPara.insert_paragraph_before("\tWithout acquiescing to propriety of the rejection, Applicant will consider submitting a terminal disclaimer, if such disclaimer is in fact needed, to overcome the double patenting rejections at a time when the claims are otherwise in condition for allowance.\n")
        DblPara.paragraph_format.line_spacing = 1.5

    if entry112 == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tRejection under 35 U.S.C. §§ 112",style="Paragraph Header")
        i += 1
        text112 = ''
        for Rej in rej_112:#Combine each rejection into one
            text112 = text112+Rej.text+' '
        Para112 = startPara.insert_paragraph_before("\tThe Examiner has rejected claims "+"under 35 U.S.C. § 112 for failing to particularly point out and distinctly claim the subject matter which the inventor regards as the invention.  Without acquiescing to the propriety of the rejection, Applicant has amended the claims to remedy the deficiencies.  Accordingly, Applicant respectfully requests withdrawal of the § 112 rejections\n")
        Para112.paragraph_format.line_spacing = 1.5

    if entry101 == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tRejections under 35 U.S.C. §§ 101",style="Paragraph Header")
        i += 1
        text101 = ''
        for Rej in rej_101:#Combine each rejection into one
            text101 = text101+Rej.text+' '
        Para101 = startPara.insert_paragraph_before("\t"+text101+"\n")
        Para101.paragraph_format.line_spacing = 1.5  

    if entry102 == True or entry103 == True:
        startPara.insert_paragraph_before("\t"+romNums[i]+"\tRejections under 35 U.S.C. §§ 102 and 103",style="Paragraph Header")
        i += 1
        text102_3 = ''
        for Rej in rej_103:#Combine each rejection into one
            text102_3 = text102_3+Rej.text+' '
        Para102_3 = startPara.insert_paragraph_before("\t"+text102_3+"\n")  
        Para102_3.paragraph_format.line_spacing = 1.5   

    #Fills Table on front page
    Template.tables[0]._cells[2].text = 'Inventor' #Inventor
    Template.tables[0]._cells[5].text = 'Examiner' #Examiner
    Template.tables[0]._cells[14].text = OA.appNum #Serial Number
    Template.tables[0]._cells[17].text = OA.artUnit #Group Art Unit
    Template.tables[0]._cells[26].text = OA.filDate #Filing Date
    Template.tables[0]._cells[29].text = OA.custNum #Customer Number
    Template.tables[0]._cells[38].text = OA.confNum #Confirmation Number
    Template.tables[0]._cells[50].text = OA.title #Title
    Template.tables[0]._cells[53].text = OA.matter #Matter Number

    #Fill small paragraph on first page
    Template.paragraphs[11].text = '\tIn response to the Non-Final Office Action dated '+OA.OAdate+', please amend the present application as set forth below and consider the remarks that follow.'

    #Fill in Header
    TempHeader = Template.sections[0].header
    TempHeader.paragraphs[0].text = 'Serial Number: '+OA.appNum+'\t\tAttorney Docket Number: '+OA.matter
    TempHeader.paragraphs[1].text = 'Response Dated: '+OA.curDate
    TempHeader.paragraphs[2].text = 'Office Action Date: '+OA.OAdate

    Template.save(TemplateFileName)
    print('End of Program.')
    time.toc()
    tkinter.messagebox.showinfo("Generation Complete.", "Template Successfully Generated.")
    exit()

Win = tkinter.Tk()
Win.title('Template Maker v2.0')
winIcon = tkinter.PhotoImage(file="C:/Users/ctroc/Documents/Python/TemplateMaker/DANDM.png",format='png')
Win.iconphoto(True,winIcon)
Win.columnconfigure(0,weight=1)
Win.rowconfigure(0,weight=1)

winFrame = GUI.Frame(Win)
winFrame.grid(column=0,row=0)

input = tkinter.StringVar()
userEntry = GUI.Entry(winFrame, width=10,textvariable=input)
userEntry.grid(column=2, row= 1)
userEntry.grid_configure(padx=5,pady=10)

entryLabel = GUI.Label(winFrame,text='Enter Application #:')
entryLabel.grid(column=1,row=1)
entryLabel.grid_configure(padx=5,pady=10)

enterButton = GUI.Button(winFrame, text='Enter', command=makeTemplate)
enterButton.grid(column=2, row=2)
enterButton.grid_configure(padx=3,pady=10)

exitButton = GUI.Button(winFrame, text='Cancel', command=exit)
exitButton.grid(column=3,row=2)
exitButton.grid_configure(padx=3,pady=10)

userEntry.focus()
Win.bind("<Return>",makeTemplate)
Win.mainloop()
